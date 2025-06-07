import logging
from logging import raiseExceptions
import docx
from typing import List, Tuple
import spire.doc
import spire.doc.FileFormat
import markdown
import tempfile
import traceback
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docxtpl import DocxTemplate
import os
import json
import uuid
from dataclasses import dataclass
from typing import Dict, Callable, Tuple, Optional
from LLMDispatch.LLMHandle.LLMWorker.Text2TextModel import TextGenerationManager
from LLMDispatch.LLMHandle.LLMWorker.Text2MindMapModel import MindMapGenerationManager
from LLMDispatch.LLMHandle.LLMWorker.Text2EChartModel import EChartGenerationManager
from LLMDispatch.LLMHandle.LLMWorker.Text2ImageModel import PictureGenerationManager
from LLMDispatch.LLMHandle.LLMWorker.Text2VideoModel import VideoGenerationManager
from LLMDispatch.LLMHandle.LLMWorker.Text2PPTModel import PptGenerationManager
from LLMDispatch.LLMHandle.LLMWorker.BoChaModel import SearchGenerationManager


@dataclass
class LLMMaster:
    def default_run_llm_task(self, task_type: str, task_model: str, task_query: str, **kwargs):
        # 注册任务调度函数，每种 task_type 映射到一个 handler 函数 返回一个可选的str
        print(
            f"task_type: {task_type}, task_model: {task_model}, task_query: '{str(task_query)[:50]}...', kwargs: {kwargs}")
        dispatcher: Dict[str, Callable[..., Any]] = {
            "default": self._handle_text_task,
            "default_stream": self._handle_text_stream_task,
            "default_thinking_stream": self._handle_text_thinking_stream_task,
            "meeting_summary": self._handle_summarizer_task,
            "meeting_summary_stream": self._handle_summarizer_stream_task,
            "meeting_summary_thinking_stream": self._handle_summarizer_thinking_stream_task,
            "content_polish_stream": self._handle_polisher_stream_task,
            "content_polish_thinking_stream": self._handle_polisher_thinking_stream_task,
            "mindmap": self._handle_mindmap_task,
            "chart": self._handle_echart_task,
            "picture": self._handle_picture_task,
            "video": self._handle_video_task,
            "formatting": self._handle_formatting_task,
            "ppt": self._handle_ppt_task,
            "outline": self._handle_outline_task,
            "search": self.__handle_search_task,
        }
        handler = dispatcher.get(task_type)
        if not handler:
            logging.error(f"不支持的 task_type: {task_type}")
            return "Unsupported task type"
        # 根据任务类型准备特定的参数
        if task_type == "formatting":
            template = kwargs.pop("template", None)
            if template is None:
                print("Formatting task requires a 'template' argument.")
                return "Formatting task requires a 'template' argument."
            return handler(task_model, task_query, template=template, **kwargs)
        elif task_type == "ppt":
            templateId = kwargs.pop("templateId", None)
            if templateId is None:
                print("PPT task requires a 'templateId' argument.")
                return "PPT task requires a 'templateId' argument."
            # 对于PPT任务，task_query被视为outline
            return handler(task_model, outline=task_query, templateId=templateId, **kwargs)
        else:
            return handler(task_model, task_query, **kwargs)

    @staticmethod
    def _handle_text_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="default", **kwargs)
        return worker.generate_text(query)

    @staticmethod
    def _handle_text_stream_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="default", **kwargs)
        return worker.generate_text_stream(query)

    @staticmethod
    def _handle_text_thinking_stream_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="default", **kwargs)
        return worker.deepthink_generate_text_stream(query)

    @staticmethod
    def _handle_summarizer_stream_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="summarizer", **kwargs)
        return worker.generate_text_stream(query)

    @staticmethod
    def _handle_summarizer_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="summarizer", **kwargs)
        return worker.generate_text(query)

    @staticmethod
    def _handle_summarizer_thinking_stream_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="summarizer", **kwargs)
        return worker.deepthink_generate_text_stream(query)

    @staticmethod
    def _handle_polisher_stream_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="polisher", **kwargs)
        return worker.generate_text_stream(query)

    @staticmethod
    def _handle_polisher_thinking_stream_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="polisher", **kwargs)
        return worker.deepthink_generate_text_stream(query)

    @staticmethod
    def _handle_outline_task(model: str, query: str, **kwargs) -> str:
        if model not in TextGenerationManager._registry:
            return f"Text model {model} not registered"
        worker = TextGenerationManager(use_api=model, role="powerpoint", **kwargs)
        return worker.generate_text(query)

    @staticmethod
    def _handle_mindmap_task(model: str, query: str, **kwargs) -> str:
        if model not in MindMapGenerationManager._registry:
            return f"MindMap model {model} not registered"
        worker = MindMapGenerationManager(use_api=model, role="mindmap")
        return worker.execute(query)

    @staticmethod
    def _handle_echart_task(model: str, query: str, **kwargs) -> str | tuple[str, str]:
        if model not in EChartGenerationManager._registry:
            return f"EChart model {model} not registered"
        worker = EChartGenerationManager(use_api=model, role="chartgen")
        return worker.execute(query)

    @staticmethod
    def _handle_picture_task(model: str, query: str, **kwargs) -> str | tuple[str, str]:
        if model not in PictureGenerationManager._registry:
            return f"Picture model {model} not registered"
        worker = PictureGenerationManager(use_api=model)
        return worker.generate_image(query, **kwargs)

    @staticmethod
    def _handle_video_task(model: str, query: str, **kwargs) -> str | tuple[str, str]:
        if model not in VideoGenerationManager._registry:
            return f"Video model {model} not registered"
        worker = VideoGenerationManager(use_api=model)
        return worker.generate_video(query, **kwargs)

    @staticmethod
    def markdown_to_docx(markdown_text, docx_file_path):
        """
        将 Markdown 文件转换为 DOCX 文件，使用 markdown 库解析，spire.doc 保存。
        通过将 HTML 写入临时文件再加载的方式。

        Args:
            markdown_text (str): 输入的 Markdown.
            docx_file_path (str): 输出的 DOCX 文件路径.
        """
        # 1. 将 Markdown 文本转换为 HTML 字符串
        try:
            # 使用 markdown 库将 Markdown 转换为 HTML 字符串
            html_content = markdown.markdown(markdown_text)
            print("Markdown 转换为 HTML 成功.")
        except Exception as e:
            print(f"将 Markdown 转换为 HTML 时发生错误: {e}")
            raise Exception(f"将 Markdown 转换为 HTML 时发生错误: {e}")

        # --- 关键修正部分 ---
        # spire.doc 不能直接从 HTML 字符串加载，需要通过文件
        temp_html_file_path = None  # 用于存储临时文件路径
        document = None  # 用于存储 spire.doc.Document 对象

        try:
            # 2. 将 HTML 内容写入一个临时 .html 文件
            # tempfile.NamedTemporaryFile 创建一个临时文件，delete=False 意味着在文件关闭后不立即删除，我们需要 spire.doc 加载后再手动删除。
            # suffix='.html' 确保文件有 .html 扩展名，spire.doc 可能依赖这个来识别格式。
            # mode='w', encoding='utf-8' 确保以文本模式和UTF-8编码写入。
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_file:
                temp_html_file_path = temp_file.name  # 获取临时文件的完整路径
                temp_file.write(html_content)  # 将 HTML 内容写入临时文件
            print(f"HTML 内容已成功写入临时文件: {temp_html_file_path}")

            # 3. 使用 spire.doc 加载临时 HTML 文件
            document = spire.doc.Document()
            # 使用 LoadFromFile 方法，指定临时文件路径和文件格式为 Html
            document.LoadFromFile(temp_html_file_path, spire.doc.FileFormat.Html)
            print("临时 HTML 文件加载到 spire.doc 文档成功.")

            # 4. 保存文档为 DOCX
            # 确保输出目录存在 (如果需要)
            output_dir = os.path.dirname(docx_file_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                print(f"创建了输出目录: {output_dir}")

            document.SaveToFile(docx_file_path, spire.doc.FileFormat.Docx)
            print(f"成功将文档保存为 DOCX 文件: {docx_file_path}")

            document = Document(docx_file_path)

            # 创建缩进距离对象
            indent_distance = Inches(0.5)
            text_to_delete_list = ['Evaluation Warning: The document was created with Spire.Doc for Python.',
                                   '```markdown']

            # 定义需要排除缩进的段落样式名称列表
            # 这些是Word中常见的标题样式名称，你可能需要根据你的文档实际使用的样式进行调整
            excluded_styles = [
                'Title',
                'Subtitle',
                'Heading 1',
                'Heading 2',
                'Heading 3',
                'Heading 4',
                'Heading 5',
                'Heading 6',
                'List Paragraph'  # 有些列表项的首行不是缩进，也可能需要排除
                # 根据你的文档实际情况添加或移除其他样式名称
            ]
            print(f"将排除以下样式进行缩进: {', '.join(excluded_styles)}")

            # 标记是否进行了修改
            made_changes = False

            # 遍历并处理文档的段落
            print("正在处理段落、删除文本和应用格式...")
            for i, paragraph in enumerate(document.paragraphs):
                original_text = paragraph.text
                original_indent = paragraph.paragraph_format.first_line_indent if paragraph.paragraph_format else None
                original_alignment = paragraph.paragraph_format.alignment if paragraph.paragraph_format else None
                current_style_name = paragraph.style.name if paragraph.style else "未知样式"

                # --- 删除文本 ---
                for text_to_delete in text_to_delete_list:
                    if text_to_delete in paragraph.text:
                        # 替换段落中的所有匹配文本为空字符串
                        # 注意：这会用一个新Run替换掉原有的Run，可能会丢失原有格式
                        paragraph.text = paragraph.text.replace(text_to_delete, "")
                        if paragraph.text != original_text:  # 检查文本是否真的发生了变化
                            made_changes = True
                            # print(f"  段落 {i+1} ({current_style_name}): 已删除文本") # 可选：打印删除信息

                # --- 设置格式 (居中或首行缩进) ---
                if paragraph.paragraph_format:  # 确保有 paragraph_format 属性
                    # 检查是否是一级标题 (Heading 1)
                    if current_style_name == 'Heading 1':
                        # 设置为居中对齐
                        if original_alignment != WD_ALIGN_PARAGRAPH.CENTER:
                            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            made_changes = True
                            # print(f"  段落 {i+1} ({current_style_name}): 已设置为居中对齐") # 可选：打印居中信息

                    # 如果不是一级标题，并且样式不在排除首行缩进的列表里
                    elif current_style_name not in excluded_styles:
                        # 设置首行缩进
                        if original_indent != indent_distance:
                            paragraph.paragraph_format.first_line_indent = indent_distance
                            made_changes = True
                            # print(f"  段落 {i+1} ({current_style_name}): 已设置首行缩进") # 可选：打印缩进信息

                    # else:
                    # print(f"  段落 {i+1} ({current_style_name}): 跳过格式设置") # 可选：打印跳过信息

            # 遍历并处理文档中的表格 (如果存在)
            print("正在处理表格中的段落、删除文本和应用格式...")
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # 遍历单元格内的段落
                        for j, paragraph in enumerate(cell.paragraphs):
                            original_text = paragraph.text
                            original_indent = paragraph.paragraph_format.first_line_indent if paragraph.paragraph_format else None
                            original_alignment = paragraph.paragraph_format.alignment if paragraph.paragraph_format else None
                            current_style_name = paragraph.style.name if paragraph.style else "未知样式"

                            # --- 删除文本 ---
                            for text_to_delete in text_to_delete_list:
                                if text_to_delete in paragraph.text:
                                    # 替换单元格段落中的所有匹配文本为空字符串
                                    # 注意：这会用一个新Run替换掉原有的Run，可能会丢失原有格式
                                    paragraph.text = paragraph.text.replace(text_to_delete, "")
                                    if paragraph.text != original_text:  # 检查文本是否真的发生了变化
                                        made_changes = True
                                        # print(f"  表格单元格中的段落 {j+1} ({current_style_name}): 已删除文本") # 可选：打印删除信息

                            # --- 设置格式 (居中或首行缩进) ---
                            if paragraph.paragraph_format:
                                # 检查是否是一级标题 (Heading 1)
                                if current_style_name == 'Heading 1':
                                    # 设置为居中对齐
                                    if original_alignment != WD_ALIGN_PARAGRAPH.CENTER:
                                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        made_changes = True
                                        # print(f"  表格单元格中的段落 {j+1} ({current_style_name}): 已设置为居中对齐") # 可选：打印居中信息

                                # 如果不是一级标题，并且样式不在排除首行缩进的列表里
                                elif current_style_name not in excluded_styles:
                                    # 设置首行缩进
                                    if original_indent != indent_distance:
                                        paragraph.paragraph_format.first_line_indent = indent_distance
                                        made_changes = True

            # 保存修改后的文档，直接覆盖原文件
            if made_changes:
                try:
                    document.save(docx_file_path)  # 直接保存到原文件路径
                    print(f"处理完成，修改已直接保存到原文件: {docx_file_path}")

                except Exception as e:
                    print(f"错误：无法保存文档，原文件可能未被修改或只部分修改 - {e}")

            else:
                print("文档中没有找到需要删除的文本，或缩进已正确设置且没有需要缩进的非标题段落，未进行保存。")

            # 保存修改后的文档
            try:
                document.save(docx_file_path)
                print(f"处理完成，修改后的文档已保存到: {docx_file_path}")
            except Exception as e:
                print(f"错误：无法保存文档 - {e}")

        except Exception as e:
            print(f"使用 spire.doc 进行转换或保存时发生错误: {e}")
            # 打印详细错误信息有助于调试
            traceback.print_exc()
            raise Exception(f"使用 spire.doc 进行转换或保存时发生错误: {e}")

        finally:
            # 5. 清理：关闭文档对象并删除临时文件
            try:
                if document is not None:
                    document.Close()  # 关闭文档对象释放资源
                    print("spire.doc 文档对象已关闭.")
            except Exception as e:
                print(e)
            # 删除临时文件，确保即使在转换过程中发生错误也能尝试删除
            if temp_html_file_path and os.path.exists(temp_html_file_path):
                try:
                    os.remove(temp_html_file_path)
                    print(f"临时文件已删除: {temp_html_file_path}")
                except Exception as e:
                    print(f"删除临时文件时发生错误: {e}")
                    raise Exception(f"删除临时文件时发生错误: {e}")

    @staticmethod
    def _handle_formatting_task(model: str, query: str, template: str, **kwargs) -> str | Tuple[str, str]:
        if model not in TextGenerationManager._registry:
            return f"Formatting model {model} not registered"
        worker = TextGenerationManager(use_api=model, role=template, **kwargs)
        text = worker.generate_text(query)
        # 生成文件名
        doc_uuid = str(uuid.uuid4())
        results_DIR = os.path.join(os.path.dirname(__file__), "..", "LLMWorker", "results")
        doc_folder = os.path.join(results_DIR, "formatting")
        os.makedirs(doc_folder, exist_ok=True)
        doc_path = os.path.join(doc_folder, f"{doc_uuid}.docx")
        if 'normal' in template:
            print(text)
            markdown_to_docx(text, doc_path)
        else:
            DOC_DIR = os.path.join(os.path.dirname(__file__), "..", "docx_template")
            tempfile = os.path.join(DOC_DIR, f"{template.replace("formatting_", "")}_template.docx")
            tpl = DocxTemplate(tempfile)
            import string
            # string.whitespace 包含了所有标准空白字符: ' \t\n\r\x0b\x0c'
            # 创建一个转换表，指定要删除 string.whitespace 中的所有字符
            translation_table = str.maketrans('', '', string.whitespace)
            # 应用转换表
            text = text.translate(translation_table)
            text = json.loads(text.replace('```json', '').replace('```', ''))
            print(text)
            tpl.render(text)
            tpl.save(doc_path)
        return doc_uuid, doc_path

    @staticmethod
    def _handle_ppt_task(model: str, outline, templateId, **kwargs):
        if model not in PptGenerationManager._registry:
            return f"PPT model {model} not registered"
        worker = PptGenerationManager(use_api=model, **kwargs)
        return worker.get_ppt_by_outline(outline, templateId)

    @staticmethod
    def __handle_search_task(model: str, query: str, **kwargs):
        if model not in SearchGenerationManager._registry:
            return f"Search model {model} not registered"
        worker = SearchGenerationManager(use_api=model, **kwargs)
        return worker.web_search(query)


if __name__ == '__main__':
    obj = LLMMaster()

    # 普通对话测试
    # task_type = 'default'
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # query = """为什么下雨过后不会一定出现彩虹？"""
    # for task_model in task_model_list:
    #     x = obj.default_run_llm_task(task_type, task_model, query)
    #     print(x)
    #     print("\n")

    # # 普通对话测试(流式)
    # task_type = 'default_stream'
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # query = """为什么下雨过后不会一定出现彩虹？"""
    # for task_model in task_model_list:
    #     x = obj.default_run_llm_task(task_type, task_model, query)
    #     for a in x:
    #         print(a, end="", flush=True)
    #     print("\n")

    # # 普通对话测试（深度思考流式）
    # task_type = 'default_thinking_stream'
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # query = """为什么下雨过后不会一定出现彩虹？"""
    # for task_model in task_model_list:
    #     combined_stream = obj.default_run_llm_task(task_type, task_model, query)
    #     print("思维链 (Streaming):")
    #     full_reasoning = []
    #     full_response = []
    #     # 遍历单个生成器
    #     for chunk in combined_stream:
    #         if chunk.type == 'reasoning':
    #             # 为了在同一行显示，我们可能需要一些技巧，或者直接分行打印
    #             print(f"{chunk.content}", end="", flush=True)
    #             full_reasoning.append(chunk.content)
    #         elif chunk.type == 'response':
    #             # 确保思考过程的输出在新的一行开始
    #             if not full_response:
    #                 print("\n--------------------")  # 在第一个响应块前打印分隔符
    #                 print("模型回复 (Streaming):")
    #
    #             print(chunk.content, end="", flush=True)
    #             full_response.append(chunk.content)
    #     print("\n")

    # # 文字总结助手（流式）
    # task_type = 'meeting_summary_stream'
    # query = """用户问为什么下雨过后不一定会有彩虹。首先，我需要回忆一下彩虹形成的基本条件。彩虹是由于阳光折射和反射在水滴中形成的，对吧？所以可能需要同时有阳光和雨滴才能形成。那可能下雨过后虽然有水滴，但如果没有阳光的话，就不会出现彩虹。不过用户的问题可能更深入一些，比如即使下雨后天空放晴，有时候还是没有彩虹。这时候可能要考虑其他因素。比如太阳的位置，必须是太阳在观察者的背后，并且角度合适，大约42度左右的仰角。如果太阳太低或者太高，可能光线无法正确折射和反射形成彩虹。还有可能水滴的大小和分布情况。如果雨后的水滴太小，比如雾气中的水滴，可能无法有效形成彩虹。或者雨后的水滴已经蒸发或散落，导致数量不足，不足以产生足够的折射效果。"""
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # for task_model in task_model_list:
    #     x = obj.default_run_llm_task(task_type, task_model, query)
    #     for a in x:
    #         print(a, end="", flush=True)
    #     print("\n")

    # # 文字总结助手（深度思考流式）
    # task_type = 'meeting_summary_thinking_stream'
    # query = """用户问为什么下雨过后不一定会有彩虹。首先，我需要回忆一下彩虹形成的基本条件。彩虹是由于阳光折射和反射在水滴中形成的，对吧？所以可能需要同时有阳光和雨滴才能形成。那可能下雨过后虽然有水滴，但如果没有阳光的话，就不会出现彩虹。不过用户的问题可能更深入一些，比如即使下雨后天空放晴，有时候还是没有彩虹。这时候可能要考虑其他因素。比如太阳的位置，必须是太阳在观察者的背后，并且角度合适，大约42度左右的仰角。如果太阳太低或者太高，可能光线无法正确折射和反射形成彩虹。还有可能水滴的大小和分布情况。如果雨后的水滴太小，比如雾气中的水滴，可能无法有效形成彩虹。或者雨后的水滴已经蒸发或散落，导致数量不足，不足以产生足够的折射效果。"""
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # for task_model in task_model_list:
    #     combined_stream = obj.default_run_llm_task(task_type, task_model, query)
    #     print("思维链 (Streaming):")
    #     full_reasoning = []
    #     full_response = []
    #     # 遍历单个生成器
    #     for chunk in combined_stream:
    #         if chunk.type == 'reasoning':
    #             # 为了在同一行显示，我们可能需要一些技巧，或者直接分行打印
    #             print(f"{chunk.content}", end="", flush=True)
    #             full_reasoning.append(chunk.content)
    #         elif chunk.type == 'response':
    #             # 确保思考过程的输出在新的一行开始
    #             if not full_response:
    #                 print("\n--------------------")  # 在第一个响应块前打印分隔符
    #                 print("模型回复 (Streaming):")
    #
    #             print(chunk.content, end="", flush=True)
    #             full_response.append(chunk.content)
    #     print("\n")

    # # 内容润色（流式）
    # task_type = 'content_polish_stream'
    # query = """为什么下雨过后不会一定出现彩虹？"""
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # for task_model in task_model_list:
    #     x = obj.default_run_llm_task(task_type, task_model, query)
    #     for a in x:
    #         print(a, end="", flush=True)
    #     print("\n")

    # # 内容润色（深度思考流式）
    # task_type = 'content_polish_thinking_stream'
    # query = """为什么下雨过后不会一定出现彩虹？"""
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # for task_model in task_model_list:
    #     combined_stream = obj.default_run_llm_task(task_type, task_model, query)
    #     print("思维链 (Streaming):")
    #     full_reasoning = []
    #     full_response = []
    #     # 遍历单个生成器
    #     for chunk in combined_stream:
    #         if chunk.type == 'reasoning':
    #             # 为了在同一行显示，我们可能需要一些技巧，或者直接分行打印
    #             print(f"{chunk.content}", end="", flush=True)
    #             full_reasoning.append(chunk.content)
    #         elif chunk.type == 'response':
    #             # 确保思考过程的输出在新的一行开始
    #             if not full_response:
    #                 print("\n--------------------")  # 在第一个响应块前打印分隔符
    #                 print("模型回复 (Streaming):")
    #
    #             print(chunk.content, end="", flush=True)
    #             full_response.append(chunk.content)
    #     print("\n")

    # # 大纲生成测试
    # task_type = 'outline'
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # query = """给我生成个大纲用于今晚上项目的汇报"""
    # for task_model in task_model_list:
    #     x = obj.default_run_llm_task(task_type, task_model, query)
    #     print(x)
    #     print("\n")

    # # 生成mindmap测试
    # task_type = 'mindmap'
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # query = """
    #         帮我画个思维导图：
    #         在香港飘扬了150多年的英国米字旗最后一次在这里降落后，接载查尔斯王子和离任港督彭定康回国的英国皇家游轮“不列颠尼亚”号驶离维多利亚港湾——这是英国撤离香港的最后时刻。
    #         英国的告别仪式是30日下午在港岛半山上的港督府（香港礼宾府）拉开序幕的。在蒙蒙细雨中，末任港督告别了这个曾居住了二十五任港督的庭院。
    #         4时30分，面色凝重的彭定康注视着港督旗帜在“日落余音”的号角声中降下旗杆。根据传统，每一位港督离任时，都举行降旗仪式。但这一次不同：永远都不会再有港督旗帜从这里升起了。4时40分，代表英国女王统治了香港五年的彭定康登上带有皇家标记的黑色“劳斯莱斯”，最后一次离开了港督府。
    #         掩映在绿树丛中的港督府于1855年建成 [6]，在以后的近一个半世纪中，包括彭定康在内的许多港督曾对其进行过大规模改建、扩建和装修。随着末代港督的离去，这座古典风格的白色建筑成为历史的陈迹。
    #         晚6时15分，象征英国管治结束的告别仪式在距离驻港英军总部不远的添马舰军营东面举行。停泊在港湾中的皇家游轮“不列颠尼亚”号和临近大厦上悬挂的巨幅紫荆花图案，恰好构成这个“日落仪式”的背景。
    #         此时，雨越下越大。查尔斯王子在雨中宣读英国女王伊丽莎白二世赠言说：“英国国旗就要降下，中国国旗将飘扬于香港上空。一百五十多年的英国管治即将告终。”
    #         7时45分，广场上灯火渐暗，开始了当天港岛上的第二次降旗仪式。一百五十六年前，一个叫爱德华·贝尔彻的英国舰长带领士兵占领了港岛，在这里升起了英国国旗；今天，另一名英国海军士兵在“威尔士亲王”军营旁的这个地方降下了米字旗。
    #         当然，最为世人瞩目的是子夜时分中英香港交接仪式上的易帜。在1997年6月30日的最后一分钟，米字旗在香港最后一次降下，英国对香港长达一个半世纪的统治宣告终结。
    #         在新的一天来临的第一分钟，五星红旗伴着《义勇军进行曲》冉冉升起，中国从此恢复对香港行使主权。与此同时，五星红旗在英军添马舰营区升起，两分钟前，“威尔士亲王”军营（中环军营） [7]移交给中国人民解放军，解放军开始接管香港防务。
    #         0时40分，刚刚参加了交接仪式的查尔斯王子和第28任港督彭定康登上“不列颠尼亚”号的甲板。在英国军舰“漆咸”号及悬挂中国国旗和香港特别行政区区旗的香港水警汽艇护卫下，将于1997年年底退役的“不列颠尼亚”号很快消失在南海的夜幕中。
    #         从1841年1月26日英国远征军第一次将米字旗插上海岛，至1997年7月1日五星红旗在香港升起，一共过去了一百五十六年五个月零四天。大英帝国从海上来，又从海上去。
    #         """
    # for task_model in task_model_list:
    #     x = obj.default_run_llm_task(task_type, task_model, query)
    #     print(x)
    #     print('\n')

    # # 生成chart测试
    # task_type = 'chart'
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # query = "今年四个季度的公司收入分别是：第一季度120万元，第二季度150万元，第三季度180万元，第四季度200万元。请绘制折线图表示收入增长情况。我们公司的网站在过去一周的访问量如下：周一有1200次访问，周二1300，周三1500，周四1700，周五2100，周六1800，周日1600。"
    # for task_model in task_model_list:
    #     x, y, = obj.default_run_llm_task(task_type, task_model, query)
    #     print(x)
    #     print(y)
    #     print("\n")

    # # 公文排版测试
    # task_type = 'formatting'
    # task_model_list = ['deepseek', 'doubao', 'qwen', 'xunfei']
    # TEXT = """
    # 在2024年到2025年期间，A公司在四个主要地区的市场表现出现了显著变化。具体来说，在2024年第一季度，北方地区的销售额达到1200万元，占全国销售额的30%，而南方地区则实现了900万元，占比22.5%。与此同时，东部和西部地区的销售额分别为1100万元和800万元，分别占据了27.5%和20%的市场份额。进入第二季度，北方地区销售增长了8%，南方增长了5%，东部持平，而西部下降了3%。到了2024年第三季度，由于新品发布，东部地区销售额暴涨20%，成为增长最快的地区；而北方和南方地区分别增长5%和4%，西部地区保持持平状态。2024年全年累计，北方地区总销售额达到5200万元，南方地区为3800万元，东部地区为4500万元，西部地区为3200万元。
    # """
    # for task_model in task_model_list:
    #     extra_params = {'template': 'formatting_letter'}
    #     x, y,  = obj.default_run_llm_task(task_type, task_model, TEXT, **extra_params)
    #     print(x)
    #     print(y)
    #     print("\n")

    # # 生成图片测试
    # task_type = 'picture'
    # task_model_list = ['qwen', 'kling', 'jimeng']
    # query = "小猫在月球上行走"
    # extra_params = {'aspect_ratio': '16:9'}
    # for task_model in task_model_list:
    #     x, y = obj.default_run_llm_task(task_type, task_model, query, **extra_params)
    #     print(x)
    #     print(y)
    #     print("\n")

    # # 生成video测试
    # task_type = 'video'
    # task_model_list = ['qwen', 'kling', 'jimeng']
    # query = "小猫在月球上行走"
    # extra_params = {'aspect_ratio': '9:16'}
    # for task_model in task_model_list:
    #     x, y = obj.default_run_llm_task(task_type, task_model, query, **extra_params)
    #     print(x)
    #     print(y)
    #     print("\n")

    # # ppt生成测试
    # task_type = 'ppt'
    # task_model = 'xunfei'
    # outline1 = {
    #                           "title": "香港回归：英国撤离的历史时刻",
    #                           "subTitle": "1997年香港主权交接全记录",
    #                           "chapters": [
    #                             {
    #                               "chapterTitle": "引言",
    #                               "chapterContents": [
    #                                 {
    #                                   "chapterTitle": "历史背景概述"
    #                                 },
    #                                 {
    #                                   "chapterTitle": "事件时间界定"
    #                                 }
    #                               ]
    #                             },
    #
    #                           ]
    #                         }
    # extra_params = {'templateId': '20240718627F1C2'}
    # x, y, z = obj.default_run_llm_task(task_type, task_model, outline1, **extra_params)
    # print(x)
    # print(y)
    # print(z)

    # # 联网搜索测试
    # task_type = 'search'
    # task_model = 'bocha'  # bocha
    # query = "小猫在月球上行走"
    # x = obj.default_run_llm_task(task_type, task_model, query)
    # print(x)
