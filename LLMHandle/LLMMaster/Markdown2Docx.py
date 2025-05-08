import os
from logging import raiseExceptions
import docx
from typing import List, Tuple
import spire.doc
import spire.doc.FileFormat
import markdown
import tempfile
import traceback
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_text, docx_file_path):
    """
    将 Markdown 文件转换为 DOCX 文件，使用 markdown 库解析，spire.doc 保存。
    通过将 HTML 写入临时文件再加载的方式。

    Args:
        markdown_text (str): 输入的 Markdown.
        docx_file_path (str): 输出的 DOCX 文件路径.
    """
    # 1. 将 Markdown 文本转换为 HTML 字符串
    try:
        # 使用 markdown 库将 Markdown 转换为 HTML 字符串
        html_content = markdown.markdown(markdown_text)
        print("Markdown 转换为 HTML 成功.")
    except Exception as e:
        print(f"将 Markdown 转换为 HTML 时发生错误: {e}")
        raise Exception(f"将 Markdown 转换为 HTML 时发生错误: {e}")

    # --- 关键修正部分 ---
    # spire.doc 不能直接从 HTML 字符串加载，需要通过文件
    temp_html_file_path = None  # 用于存储临时文件路径
    document = None  # 用于存储 spire.doc.Document 对象

    try:
        # 2. 将 HTML 内容写入一个临时 .html 文件
        # tempfile.NamedTemporaryFile 创建一个临时文件，delete=False 意味着在文件关闭后不立即删除，我们需要 spire.doc 加载后再手动删除。
        # suffix='.html' 确保文件有 .html 扩展名，spire.doc 可能依赖这个来识别格式。
        # mode='w', encoding='utf-8' 确保以文本模式和UTF-8编码写入。
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_file:
            temp_html_file_path = temp_file.name  # 获取临时文件的完整路径
            temp_file.write(html_content)  # 将 HTML 内容写入临时文件
        print(f"HTML 内容已成功写入临时文件: {temp_html_file_path}")

        # 3. 使用 spire.doc 加载临时 HTML 文件
        document = spire.doc.Document()
        # 使用 LoadFromFile 方法，指定临时文件路径和文件格式为 Html
        document.LoadFromFile(temp_html_file_path, spire.doc.FileFormat.Html)
        print("临时 HTML 文件加载到 spire.doc 文档成功.")

        # 4. 保存文档为 DOCX
        # 确保输出目录存在 (如果需要)
        output_dir = os.path.dirname(docx_file_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"创建了输出目录: {output_dir}")

        document.SaveToFile(docx_file_path, spire.doc.FileFormat.Docx)
        print(f"成功将文档保存为 DOCX 文件: {docx_file_path}")

        document = Document(docx_file_path)

        # 创建缩进距离对象
        indent_distance = Inches(0.5)
        text_to_delete_list = ['Evaluation Warning: The document was created with Spire.Doc for Python.', '```markdown']

        # 定义需要排除缩进的段落样式名称列表
        # 这些是Word中常见的标题样式名称，你可能需要根据你的文档实际使用的样式进行调整
        excluded_styles = [
            'Title',
            'Subtitle',
            'Heading 1',
            'Heading 2',
            'Heading 3',
            'Heading 4',
            'Heading 5',
            'Heading 6',
            'List Paragraph'  # 有些列表项的首行不是缩进，也可能需要排除
            # 根据你的文档实际情况添加或移除其他样式名称
        ]
        print(f"将排除以下样式进行缩进: {', '.join(excluded_styles)}")

        # 标记是否进行了修改
        made_changes = False

        # 遍历并处理文档的段落
        print("正在处理段落、删除文本和应用格式...")
        for i, paragraph in enumerate(document.paragraphs):
            original_text = paragraph.text
            original_indent = paragraph.paragraph_format.first_line_indent if paragraph.paragraph_format else None
            original_alignment = paragraph.paragraph_format.alignment if paragraph.paragraph_format else None
            current_style_name = paragraph.style.name if paragraph.style else "未知样式"

            # --- 删除文本 ---
            for text_to_delete in text_to_delete_list:
                if text_to_delete in paragraph.text:
                    # 替换段落中的所有匹配文本为空字符串
                    # 注意：这会用一个新Run替换掉原有的Run，可能会丢失原有格式
                    paragraph.text = paragraph.text.replace(text_to_delete, "")
                    if paragraph.text != original_text:  # 检查文本是否真的发生了变化
                        made_changes = True
                        # print(f"  段落 {i+1} ({current_style_name}): 已删除文本") # 可选：打印删除信息

            # --- 设置格式 (居中或首行缩进) ---
            if paragraph.paragraph_format:  # 确保有 paragraph_format 属性
                # 检查是否是一级标题 (Heading 1)
                if current_style_name == 'Heading 1':
                    # 设置为居中对齐
                    if original_alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        made_changes = True
                        # print(f"  段落 {i+1} ({current_style_name}): 已设置为居中对齐") # 可选：打印居中信息

                # 如果不是一级标题，并且样式不在排除首行缩进的列表里
                elif current_style_name not in excluded_styles:
                    # 设置首行缩进
                    if original_indent != indent_distance:
                        paragraph.paragraph_format.first_line_indent = indent_distance
                        made_changes = True
                        # print(f"  段落 {i+1} ({current_style_name}): 已设置首行缩进") # 可选：打印缩进信息

                # else:
                # print(f"  段落 {i+1} ({current_style_name}): 跳过格式设置") # 可选：打印跳过信息

        # 遍历并处理文档中的表格 (如果存在)
        print("正在处理表格中的段落、删除文本和应用格式...")
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 遍历单元格内的段落
                    for j, paragraph in enumerate(cell.paragraphs):
                        original_text = paragraph.text
                        original_indent = paragraph.paragraph_format.first_line_indent if paragraph.paragraph_format else None
                        original_alignment = paragraph.paragraph_format.alignment if paragraph.paragraph_format else None
                        current_style_name = paragraph.style.name if paragraph.style else "未知样式"

                        # --- 删除文本 ---
                        for text_to_delete in text_to_delete_list:
                            if text_to_delete in paragraph.text:
                                # 替换单元格段落中的所有匹配文本为空字符串
                                # 注意：这会用一个新Run替换掉原有的Run，可能会丢失原有格式
                                paragraph.text = paragraph.text.replace(text_to_delete, "")
                                if paragraph.text != original_text:  # 检查文本是否真的发生了变化
                                    made_changes = True
                                    # print(f"  表格单元格中的段落 {j+1} ({current_style_name}): 已删除文本") # 可选：打印删除信息

                        # --- 设置格式 (居中或首行缩进) ---
                        if paragraph.paragraph_format:
                            # 检查是否是一级标题 (Heading 1)
                            if current_style_name == 'Heading 1':
                                # 设置为居中对齐
                                if original_alignment != WD_ALIGN_PARAGRAPH.CENTER:
                                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    made_changes = True
                                    # print(f"  表格单元格中的段落 {j+1} ({current_style_name}): 已设置为居中对齐") # 可选：打印居中信息

                            # 如果不是一级标题，并且样式不在排除首行缩进的列表里
                            elif current_style_name not in excluded_styles:
                                # 设置首行缩进
                                if original_indent != indent_distance:
                                    paragraph.paragraph_format.first_line_indent = indent_distance
                                    made_changes = True

        # 保存修改后的文档，直接覆盖原文件
        if made_changes:
            try:
                document.save(docx_file_path)  # 直接保存到原文件路径
                print(f"处理完成，修改已直接保存到原文件: {docx_file_path}")

            except Exception as e:
                print(f"错误：无法保存文档，原文件可能未被修改或只部分修改 - {e}")

        else:
            print("文档中没有找到需要删除的文本，或缩进已正确设置且没有需要缩进的非标题段落，未进行保存。")

        # 保存修改后的文档
        try:
            document.save(docx_file_path)
            print(f"处理完成，修改后的文档已保存到: {docx_file_path}")
        except Exception as e:
            print(f"错误：无法保存文档 - {e}")

    except Exception as e:
        print(f"使用 spire.doc 进行转换或保存时发生错误: {e}")
        # 打印详细错误信息有助于调试
        traceback.print_exc()
        raise Exception(f"使用 spire.doc 进行转换或保存时发生错误: {e}")

    finally:
        # 5. 清理：关闭文档对象并删除临时文件
        try:
            if document is not None:
                document.Close()  # 关闭文档对象释放资源
                print("spire.doc 文档对象已关闭.")
        except Exception as e:
            print(e)
        # 删除临时文件，确保即使在转换过程中发生错误也能尝试删除
        if temp_html_file_path and os.path.exists(temp_html_file_path):
            try:
                os.remove(temp_html_file_path)
                print(f"临时文件已删除: {temp_html_file_path}")
            except Exception as e:
                print(f"删除临时文件时发生错误: {e}")
                raise Exception(f"删除临时文件时发生错误: {e}")
